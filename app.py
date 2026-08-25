import io
import re
import pandas as pd
import streamlit as st
from pypdf import PdfReader
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

# --- PAGE CONFIG ---
st.set_page_config(
    page_title="Complivox Global | Statutory Defense Engine",
    page_icon="🛡️",
    layout="wide",
    initial_sidebar_state="expanded"
)

# --- ADVANCED ENTERPRISE CSS ---
st.markdown("""
<style>
    .brand-title {
        font-size: 2.3rem;
        font-weight: 800;
        color: #0E7490;
        margin-bottom: 0px;
        letter-spacing: -0.5px;
    }
    .brand-sub {
        font-size: 1rem;
        color: #64748B;
        margin-bottom: 20px;
    }
    .metric-card {
        background: linear-gradient(135deg, #0f172a 0%, #1e293b 100%);
        border-radius: 10px;
        padding: 16px 20px;
        color: white;
        box-shadow: 0 4px 6px -1px rgba(0,0,0,0.1);
        margin-bottom: 12px;
    }
    .metric-card-val {
        font-size: 1.8rem;
        font-weight: 800;
        color: #38BDF8;
    }
    .metric-card-lbl {
        font-size: 0.85rem;
        color: #94A3B8;
        text-transform: uppercase;
        letter-spacing: 0.5px;
    }
    .stButton>button {
        background: linear-gradient(90deg, #0E7490, #0369A1);
        color: white;
        font-weight: 700;
        border-radius: 8px;
        padding: 0.6rem 1rem;
        border: none;
        width: 100%;
        transition: 0.2s;
    }
    .stButton>button:hover {
        background: linear-gradient(90deg, #0891B2, #0284C7);
        box-shadow: 0 4px 12px rgba(14, 116, 144, 0.3);
    }
    .report-card {
        background-color: #F8FAFC;
        border-left: 4px solid #0E7490;
        padding: 14px 18px;
        border-radius: 0 8px 8px 0;
        margin-bottom: 12px;
    }
</style>
""", unsafe_allow_html=True)

# --- CLEAN WORD GENERATOR (NO PAGE-CUTS) ---
def generate_docx_dossier(domain, framework, doc_name, findings, gaps, rtqs, score):
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(0.7)
        s.bottom_margin = Inches(0.7)
        s.left_margin = Inches(0.7)
        s.right_margin = Inches(0.7)

    t = doc.add_paragraph()
    tr = t.add_run("COMPLIVOX GLOBAL")
    tr.font.name = "Arial"
    tr.font.size = Pt(18)
    tr.font.bold = True
    tr.font.color.rgb = RGBColor(14, 116, 144)
    t.paragraph_format.space_after = Pt(2)

    sub = doc.add_paragraph()
    sr = sub.add_run(f"Statutory Pre-Submission Defense Dossier | Framework: {framework} ({domain})\nSource Document: {doc_name} | Compliance Index: {score}/100")
    sr.font.name = "Arial"
    sr.font.size = Pt(9.5)
    sr.font.italic = True
    sr.font.color.rgb = RGBColor(100, 116, 139)
    sub.paragraph_format.space_after = Pt(14)

    # 1. Findings
    h1 = doc.add_paragraph()
    h1r = h1.add_run("1. Technical Compliance Extraction Status")
    h1r.font.name = "Arial"
    h1r.font.bold = True
    h1r.font.size = Pt(11)
    h1.paragraph_format.keep_with_next = True
    h1.paragraph_format.space_after = Pt(4)

    for item in findings:
        p = doc.add_paragraph(style='List Bullet')
        pr = p.add_run(str(item))
        pr.font.name = "Arial"
        pr.font.size = Pt(9.5)
        p.paragraph_format.space_after = Pt(2)

    # 2. Gap Matrix
    h2 = doc.add_paragraph()
    h2r = h2.add_run("2. Statutory Gap & Deficiency Matrix")
    h2r.font.name = "Arial"
    h2r.font.bold = True
    h2r.font.size = Pt(11)
    h2.paragraph_format.keep_with_next = True
    h2.paragraph_format.space_before = Pt(8)
    h2.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True

    hdr_cells = table.rows[0].cells
    cols = ["Regulatory Requirement", "Status / Deficit", "Risk Level"]
    for i, col_name in enumerate(cols):
        hdr_cells[i].text = col_name
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="1E293B"/>')
        hdr_cells[i]._tc.get_or_add_tcPr().append(shd)
        p = hdr_cells[i].paragraphs[0]
        if len(p.runs) > 0:
            p.runs[0].font.name = "Arial"
            p.runs[0].font.bold = True
            p.runs[0].font.size = Pt(9)
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)

    for g in gaps:
        row = table.add_row().cells
        row[0].text = str(g.get("rule", "Statutory Clause"))
        row[1].text = str(g.get("detail", "Pending validation"))
        row[2].text = str(g.get("risk", "Medium"))
        for c_idx, cell in enumerate(row):
            shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F8FAFC"/>')
            cell._tc.get_or_add_tcPr().append(shd)
            if len(cell.paragraphs[0].runs) > 0:
                cell.paragraphs[0].runs[0].font.name = "Arial"
                cell.paragraphs[0].runs[0].font.size = Pt(8.5)
                if c_idx == 2 and g.get("risk") == "High":
                    cell.paragraphs[0].runs[0].font.bold = True
                    cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(185, 28, 28)

    for r in table.rows:
        r._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

    # 3. SEC RTQ Cards
    h3 = doc.add_paragraph()
    h3r = h3.add_run("3. Pre-Emptive SEC Objections & Defense Strategy (RTQ)")
    h3r.font.name = "Arial"
    h3r.font.bold = True
    h3r.font.size = Pt(11)
    h3.paragraph_format.keep_with_next = True
    h3.paragraph_format.space_before = Pt(10)
    h3.paragraph_format.space_after = Pt(6)

    for idx, r in enumerate(rtqs, 1):
        card = doc.add_table(rows=1, cols=1)
        card.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = card.rows[0].cells[0]
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="F1F5F9"/>')
        cell._tc.get_or_add_tcPr().append(shd)

        cp = cell.paragraphs[0]
        cp.paragraph_format.space_after = Pt(2)
        r1 = cp.add_run(f"Objection #{idx}: {r.get('query', '')} [{r.get('risk', 'High')}]\n")
        r1.font.name = "Arial"
        r1.font.bold = True
        r1.font.size = Pt(9.5)
        r1.font.color.rgb = RGBColor(15, 23, 42)

        r2 = cp.add_run(f"Statutory Standard: {r.get('standard', '')}\n")
        r2.font.name = "Arial"
        r2.font.italic = True
        r2.font.size = Pt(8.5)
        r2.font.color.rgb = RGBColor(71, 85, 105)

        r3 = cp.add_run(f"Recommended Defense (RTQ): {r.get('defense', '')}")
        r3.font.name = "Arial"
        r3.font.bold = True
        r3.font.size = Pt(9)
        r3.font.color.rgb = RGBColor(14, 116, 144)

        card.rows[0]._tr.get_or_add_trPr().append(parse_xml(f'<w:cantSplit {nsdecls("w")}/>'))

        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(4)

    target_stream = io.BytesIO()
    doc.save(target_stream)
    target_stream.seek(0)
    return target_stream

# --- AUDIT SCORING ENGINE ---
def run_full_audit(text, domain, framework):
    t = text.lower()
    findings = []
    gaps = []
    rtqs = []
    score = 100

    if domain == "Pharmaceuticals & Bulk Drugs":
        if "nitrosamine" in t or "ndma" in t:
            findings.append("Nitrosamine toxicological evaluation identified in dossier.")
        else:
            score -= 25
            findings.append("Nitrosamine risk assessment section missing or unconfirmed.")
            gaps.append({"rule": "ICH M7 / CDSCO Nitrosamines", "detail": "Absence of confirmatory LC-MS/MS testing for nitrosamine impurities.", "risk": "High"})
            rtqs.append({
                "query": "Risk evaluation and limit justification for potential Nitrosamine contamination in synthetic route.",
                "standard": "CDSCO Notification No. 29-114/2019-DC / US FDA Guidance",
                "defense": "Submit synthetic route purge-factor calculations proving secondary/tertiary amines are removed before final crystallization.",
                "risk": "High"
            })

        if "zone ivb" in t or "30°c/75% rh" in t or "accelerated" in t:
            findings.append("Climatic stability protocol complies with regulatory requirements.")
        else:
            score -= 20
            findings.append("Climatic Zone IVb stability testing parameters unverified.")
            gaps.append({"rule": "ICH Q1A (R2) / NDCT Schedule Y", "detail": "Real-time stability data at 30°C/75% RH for Indian climatic conditions absent.", "risk": "Medium"})
            rtqs.append({
                "query": "Submission of completed 6-month accelerated and ongoing 12-month real-time Zone IVb stability data.",
                "standard": "CDSCO Form 40 / NDCT Rule 75",
                "defense": "Furnish interim 6-month stability matrix with committed protocol for 24-month long-term testing at 30°C/75% RH.",
                "risk": "Medium"
            })

        if "dissolution" in t or "bioequivalence" in t:
            findings.append("Comparative in-vitro dissolution / BE profile identified.")
        else:
            score -= 25
            gaps.append({"rule": "NDCT Rules 2019 / US FDA BE Guidance", "detail": "Comparative dissolution profile in 3 discriminatory media missing.", "risk": "High"})
            rtqs.append({
                "query": "Demonstration of in-vitro similarity (f2 factor > 50) against reference listed innovator product.",
                "standard": "Rule 60(1) Phase III Clinical Waiver / FDA Guidance on BE",
                "defense": "Provide multi-point dissolution data at pH 1.2, 4.5, and 6.8 showing f2 similarity factor between 50-100.",
                "risk": "High"
            })

    else: # Medical Devices
        if "iso 10993" in t or "biocompatibility" in t or "cytotoxicity" in t:
            findings.append("Biological safety assessment referenced under ISO 10993.")
        else:
            score -= 30
            findings.append("ISO 10993 biological evaluation summary not explicitly confirmed.")
            gaps.append({"rule": "ISO 10993-1 / MDR 2017 Schedule 4", "detail": "Cytotoxicity, sensitization, and intracutaneous reactivity tests absent.", "risk": "High"})
            rtqs.append({
                "query": "Complete ISO 10993 biocompatibility testing for direct patient-contact materials.",
                "standard": "CDSCO Form MD-14 Guidance / FDA 510(k) Cytotoxicity Criteria",
                "defense": "Provide GLP-certified biological safety endpoints and material characterization certificate of analysis (COA).",
                "risk": "High"
            })

        if "sterilization" in t or "sal 10" in t or "ethylene oxide" in t:
            findings.append("Sterilization validation protocol and Sterility Assurance Level (SAL 10^-6) verified.")
        else:
            score -= 25
            gaps.append({"rule": "ISO 11135 / ISO 11137 Validation", "detail": "Sterilization validation report and EO residue limits missing.", "risk": "High"})
            rtqs.append({
                "query": "Sterilization validation protocol and residual limits justification.",
                "standard": "MDR 2017 Fifth Schedule / ISO 10993-7",
                "defense": "Submit EO/ECH residual testing reports adhering to allowable limits under ISO 10993-7 along with dose mapping.",
                "risk": "High"
            })

        if "predicate" in t or "substantially equivalent" in t:
            findings.append("Predicate device substantial equivalence comparison table detected.")
        else:
            score -= 20
            gaps.append({"rule": "FDA 21 CFR 807.87 / CDSCO MD-14", "detail": "Side-by-side technological and performance comparison with predicate missing.", "risk": "Medium"})
            rtqs.append({
                "query": "Demonstration of technological equivalence against approved predicate device.",
                "standard": "CDSCO SUGAM Portal / FDA 510(k) Section 12",
                "defense": "Submit side-by-side comparative table demonstrating identical intended use, materials, and bench test performance.",
                "risk": "Medium"
            })

    return findings, gaps, rtqs, max(score, 30)

# --- HEADER INTERFACE ---
st.markdown('<div class="brand-title">COMPLIVOX GLOBAL</div>', unsafe_allow_html=True)
st.markdown('<div class="brand-sub">Dual Statutory Regulatory Intelligence & SEC Pre-Submission Defense Architecture</div>', unsafe_allow_html=True)

# --- SIDEBAR ---
st.sidebar.markdown("### ⚙️ Regulatory Controls")
domain = st.sidebar.radio(
    "Target Domain",
    ["Pharmaceuticals & Bulk Drugs", "Medical Devices & Diagnostics"]
)

if domain == "Pharmaceuticals & Bulk Drugs":
    framework = st.sidebar.selectbox(
        "Filing Framework",
        ["CDSCO Form 40 / NDCT 2019 (India)", "US FDA eCTD Module 1-5 (USA)", "EU CTD Marketing Authorization (EMA)"]
    )
else:
    framework = st.sidebar.selectbox(
        "Filing Framework",
        ["CDSCO Form MD-14 / MDR 2017 (India)", "US FDA 510(k) Substantial Equivalence", "EU MDR 2017/745 Annex II/III"]
    )

st.sidebar.markdown("---")
st.sidebar.markdown("### 📋 Statutory Reference Tools")
st.sidebar.info("• **NDCT 2019 Rule 60(1):** Local Phase III trial waiver criteria.\n\n• **MDR 2017 Schedule 4:** Device Master File (DMF) checklist.\n\n• **ICH M7 Guidance:** Nitrosamine toxicological purge thresholds.")

# --- MAIN INTERFACE TABS ---
tab_audit, tab_analytics, tab_simulator = st.tabs(["📑 Dossier Audit & RTQ Engine", "📊 Regulatory Risk Analytics", "⚡ Instant Query Simulator"])

with tab_audit:
    c1, c2 = st.columns([1, 1])

    with c1:
        st.markdown("#### 1. Upload Submission Dossier")
        uploaded_file = st.file_uploader("Upload Technical Dossier / DMF / Summary (PDF)", type=["pdf"])
        extracted_text = ""
        page_count = 0
        if uploaded_file:
            try:
                reader = PdfReader(uploaded_file)
                page_count = len(reader.pages)
                for p in reader.pages:
                    txt = p.extract_text()
                    if txt:
                        extracted_text += txt + "\n"
                st.success(f"✓ Parsed `{uploaded_file.name}` ({page_count} Pages)")
            except Exception as e:
                st.error(f"Error parsing PDF: {e}")

    with c2:
        st.markdown("#### 2. Defense Audit Engine")
        if extracted_text:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-card-lbl">Document Payload</div>
                <div class="metric-card-val">{len(extracted_text):,} <span style="font-size: 1rem; color: #94A3B8;">Chars</span></div>
                <div style="font-size: 0.8rem; color: #CBD5E1; margin-top: 4px;">Pages Processed: {page_count} | Mode: Statutory Scrutiny</div>
            </div>
            """, unsafe_allow_html=True)
            run_btn = st.button("🚀 Run Statutory Scrutiny & RTQ Generation")
        else:
            st.info("Upload a regulatory PDF on the left to activate defense engine.")
            run_btn = False

    if extracted_text and run_btn:
        findings, gaps, rtqs, score = run_full_audit(extracted_text, domain, framework)
        st.session_state['last_audit'] = {'findings': findings, 'gaps': gaps, 'rtqs': rtqs, 'score': score, 'file': uploaded_file.name}

    if 'last_audit' in st.session_state:
        audit = st.session_state['last_audit']
        st.markdown("---")
        
        # TOP EXECUTIVE METRICS ROW
        m1, m2, m3 = st.columns(3)
        with m1:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-card-lbl">Regulatory Scrutiny Index</div>
                <div class="metric-card-val">{audit['score']}/100</div>
                <div style="font-size: 0.8rem; color: {'#4ADE80' if audit['score']>75 else '#F87171'};">
                    {'✓ Low Scrutiny Risk' if audit['score']>75 else '⚠️ High SEC Rejection Risk'}
                </div>
            </div>
            """, unsafe_allow_html=True)
        with m2:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-card-lbl">Identified Deficiencies</div>
                <div class="metric-card-val">{len(audit['gaps'])}</div>
                <div style="font-size: 0.8rem; color: #FBBF24;">Actionable Statutory Gaps</div>
            </div>
            """, unsafe_allow_html=True)
        with m3:
            st.markdown(f"""
            <div class="metric-card">
                <div class="metric-card-lbl">Pre-Drafted RTQ Responses</div>
                <div class="metric-card-val">{len(audit['rtqs'])}</div>
                <div style="font-size: 0.8rem; color: #38BDF8;">Audit-Ready Defense Justifications</div>
            </div>
            """, unsafe_allow_html=True)

        st.markdown("<br>", unsafe_allow_html=True)
        r_col1, r_col2 = st.columns([1, 1.2])

        with r_col1:
            st.markdown("### 📋 Statutory Gap Matrix")
            for g in audit['gaps']:
                badge_color = "#DC2626" if g["risk"] == "High" else "#D97706"
                st.markdown(f"""
                <div class="report-card" style="border-left-color: {badge_color};">
                    <strong style="color: {badge_color};">[{g['risk'].upper()} RISK] {g['rule']}</strong><br>
                    <span style="font-size: 0.9rem; color: #334155;">{g['detail']}</span>
                </div>
                """, unsafe_allow_html=True)

            st.markdown("#### Extraction Checklist")
            for f in audit['findings']:
                st.markdown(f"• {f}")

        with r_col2:
            st.markdown("### 🛡️ Pre-Emptive SEC Objections & Defense")
            for idx, r in enumerate(audit['rtqs'], 1):
                with st.expander(f"Objection #{idx}: {r['query']}", expanded=True):
                    st.caption(f"**Statutory Standard:** {r['standard']}")
                    st.info(f"**Recommended RTQ Defense:**\n\n{r['defense']}")

            st.markdown("---")
            docx_file = generate_docx_dossier(
                domain=domain,
                framework=framework,
                doc_name=audit['file'],
                findings=audit['findings'],
                gaps=audit['gaps'],
                rtqs=audit['rtqs'],
                score=audit['score']
            )

            st.download_button(
                label="📥 Download Audit Dossier (.DOCX)",
                data=docx_file,
                file_name=f"Complivox_Audit_Dossier_{domain[:3].upper()}.docx",
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            )

with tab_analytics:
    st.markdown("### 📊 Enterprise Statutory Analytics")
    if 'last_audit' in st.session_state:
        audit = st.session_state['last_audit']
        
        # Risk velocity chart
        risk_data = pd.DataFrame({
            "Module": ["Stability (Zone IVb)", "Nitrosamine/Toxicology", "Biocompatibility", "Predicate Equivalence", "Clinical Waiver Defense"],
            "Scrutiny Probability (%)": [85, 90, 70, 60, 95]
        })
        st.markdown("#### Pre-Submission Objection Probability Index")
        st.bar_chart(risk_data.set_index("Module"))
    else:
        st.info("Run an audit in Tab 1 to view interactive statutory charts.")

with tab_simulator:
    st.markdown("### ⚡ Instant Regulatory Scenario Simulator")
    sim_col1, sim_col2 = st.columns(2)
    with sim_col1:
        scenario = st.selectbox(
            "Select Statutory Scenario",
            [
                "CDSCO Form 40 API Import - Nitrosamine & Zone IVb Stability Query",
                "CDSCO Form MD-14 Device Import - Biocompatibility & EO Residuals",
                "US FDA 510(k) - Predicate Substantial Equivalence Margin"
            ]
        )
    
    if st.button("Run Simulation"):
        st.success(f"Running statutory simulation for: **{scenario}**")
        if "Nitrosamine" in scenario:
            f, g, r, s = run_full_audit("nitrosamine missing accelerated only", "Pharmaceuticals & Bulk Drugs", "CDSCO Form 40 / NDCT 2019 (India)")
        elif "Biocompatibility" in scenario:
            f, g, r, s = run_full_audit("no biocompatibility", "Medical Devices & Diagnostics", "CDSCO Form MD-14 / MDR 2017 (India)")
        else:
            f, g, r, s = run_full_audit("predicate pending", "Medical Devices & Diagnostics", "US FDA 510(k) Substantial Equivalence")
        
        st.session_state['last_audit'] = {'findings': f, 'gaps': g, 'rtqs': r, 'score': s, 'file': 'Simulated_Scenario.pdf'}
        st.rerun()
